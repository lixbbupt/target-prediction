# -*- coding: utf-8 -*-
"""
Created on Thu Jul 19 13:11:40 2018

dataset: three xlsx files
    file 1: HT_20171031_MJZ1_Extension(20180312).xlsx
    file 2: HT_data_merge_aminoAcid_raw.xlsx
    file 3: HT_data_Pirche.xlsx

#### stage 0:

task 1: read xlsx file into dataframe using python pandas and extract the following columns respectively,
    file 1: ID + non-days-related columns
    file 2: ID + columns from 'AM_A other mismatch' to the end 
    file 3: ID + PIRCHE_I_num + PIRCHE_II_num
    
task 2:
    merge the extracted data into one dataframe based on ID, then save the dataframe to csv file

task 3:
    find out categorical features and continuous features
    
task 4:
    select a few features and do some statistics for them, such as missing rate, mean, median.
    plot for continuous features

#### stage 1 (in about 8-12 days): pre-modeling

task 5: data preprocessing: 
    a) make clear of target variables and feature variables
    b) handle missing values for features (methods survey and implementation)
    c) try one-hot encoding for categorical features

task 6: feature engineering-1: feature generating
    a) make combinations for features (and/or/xor/+/-/x/... etc.)

task 7: feature engineering-2: feature selection by information value
    a) information value calculation (search GitHub)
    b) try feature selection using information value(要比较用woe 填充和不用的区别，等建模后来测一下)

task 8: feature engineering-3: feature selection using scikit-learn/skfeature packages
    a) try feature selection methods in scikit-learn packages
    b) try feature selection methods in skfeature packages(看参考)

#### stage 2 (in about 10-15 days): modeling for 

task 9: Logistic Regression model 
task 10: XGBoost, lightGBM, CatBoost model (parameters tuning, visualization)
task 11: model ensembling

#### stage 3 (in about 4-6 days): analysis

task 12: analyze the impact on modeling for the four groups of features,  ie., HLA-related features, KIR-related features, PIRCHE-related features, amino-acid-related features
task 13: a final report


@author: xbli
"""

import pandas as pd
import numpy as np
#from numpy import array
#from numpy import argmax
#from sklearn.preprocessing import LabelEncoder
#from sklearn.preprocessing import OneHotEncoder
#import os
#from IPython.display import display
from operator import itemgetter

#import os
import numpy as np
import woe.feature_process as fp
#import woe.GridSearch as gs
import sklearn.linear_model as lm
from sklearn.ensemble import RandomForestClassifier
import xgboost as xgb

from sklearn import preprocessing

import matplotlib.pyplot as plt
from sklearn.svm import SVC
from sklearn.model_selection import StratifiedKFold
from sklearn.feature_selection import RFECV
from sklearn.datasets import make_classification

from sklearn.model_selection import train_test_split
from sklearn.metrics import roc_auc_score
from datetime import datetime

import sys
sys.path.insert(0,'/Users/DihuaDuolan/Desktop/BFR')
from task1To4.task_0719_0724 import missing_value_selection
from task1To4.task_0719_0724 import cat_cot_dis_data
from task1To4.task_0719_0724 import seperate_patient_donor_blood_type

#logistic regression
import statsmodels.api as sm
from sklearn import metrics
from sklearn.model_selection import cross_val_score
from sklearn.model_selection import GridSearchCV
from sklearn.metrics import classification_report


from docx import Document
from statistics import mean
import pickle
import seaborn as sns # data visualization library 


''' 建造对于categorical feature 的 one hot encoding '''
def cat_feature_to_discreate(column, df_data1):

    # 对每个feature 进行名字拆分
    df_temp = pd.get_dummies(df_data1[column])
    df_temp.columns = [column+'_'+ item for item in df_temp.columns]
    # append new cols to data frame and  delete original column 
    return df_temp


'''将 DIS 和 aminoAcid 的 特征拆分成离散. eg.aminoAcid_L'''
def one_hot_encoding(df_data3, cols_one_hot):
    df_temp_list = []
    # one hot encoding for 'DIS' and aminoAcid
    for col in cols_one_hot:
        df_temp_list.append(cat_feature_to_discreate(col, df_data3))
        
    df_data3 = df_data3.join(df_temp_list).drop(cols_one_hot, axis=1)
    return df_data3

''' 缺失值处理, 类别型categorical: 填写一个没有出现的值(3)；数值型：填写均值 '''
def fill_missing_value(df_data4, col_cat, col_num_cot, col_num_dis, fill_num):
    
    # ============================= cot data, 填写均值
    for i in col_num_cot:
        df_data4[i].fillna(df_data4[i].mean(), inplace=True)
    
    # ================= discrete feature(in cat) 'hla', 填写均值
    for i in col_num_dis:
        df_data4[i].fillna(df_data4[i].median(), inplace=True)
    
    # ====== cat data: 'hla' not in i and i not in delete_col 填写从未出现过的数字
    for i in col_cat:
        df_data4[i] = df_data4[i].fillna(fill_num)
    
    return df_data4

''' 建造给information value 用的config csv. 内容包括 data type & featue & target等等'''
def config_csv(hospital_name, col_name, new_path, col_cot, features_variables_col, target):
    
    # continuous & target -> int64, category (discrete) -> object
    col_dtype = ['int64' if (i in col_cot) or (i == target) else 'object' for i in col_name]
    
    # continuous -> 1, category (discrete) -> 0
    col_is_tobe_bin = [1 if i in col_cot else 0 for i in col_name]
    
    # all the things except target and ID and dates
    col_is_candidate = [1 if (i in features_variables_col) and (i != 'ID') else 0 for i in col_name]
    col_is_model_feature = col_is_candidate
    
    # write to excel and csv
    df_data = {'var_name':col_name[0:-1]+['target'],
         'var_dtype':col_dtype,
         'is_tobe_bin':col_is_tobe_bin,
         'is_candidate': col_is_candidate,
         'is_modelfeature': col_is_model_feature}
    df = pd.DataFrame.from_dict(df_data)
    
    df.to_excel(new_path + 'config_excel_'  + target + hospital_name + '.xlsx', index=False, encoding= 'utf-8')
    
    df.to_csv(new_path + 'config_csv_'  + target + hospital_name + '.csv', index=False, encoding= 'utf-8')
    return df


def get_info_value_dic(feature_detail_path):
    info_value_df = pd.read_csv(feature_detail_path)
    info_value_dic = info_value_df.groupby('var_name')['iv'].apply(lambda x: x.mean()).to_dict()
    
    # return the sorted info_value dic
    sorted_info_value_dic = sorted(info_value_dic.items(), key=lambda x: x[1])

    
    # return 一个dic {feature name， infromation value}
    return sorted_info_value_dic

def information_value(hospital_name, df_data5, new_path, col_num_cot, features_variables_col, target):
    
    # 第一步： create a csv file
    col_name = list(features_variables_col) + [target]
     # for quick debug only ====== 选取一小段来检测
    #col_name = list(features_variables_col[0:200] + ['CGVHD']) 
    
    
    config_csv(hospital_name, col_name, new_path, col_num_cot, features_variables_col, target)
    
    # 第二步： 用git hub 上面贡献的代码来算 infromation value，output 一个dictionary
    
    df_data_woe = df_data5[col_name]
    df_data_woe.columns = col_name[0:-1]+['target']
        
    config_path =  new_path +'config_csv_'  + target + hospital_name + '.csv'
    data_path = new_path +'data_woe_' + target + hospital_name + '.csv'
    feature_detail_path = new_path +'features_detail_'  + target + hospital_name +'.csv'
    rst_pkl_path = new_path +'woe_rule_' + target + hospital_name +'.pkl'
    
    df_data_woe.to_csv(data_path,index=False, encoding='utf-8')

    # train woe rule
    feature_detail,rst = fp.process_train_woe(infile_path=data_path,outfile_path=feature_detail_path,rst_path=rst_pkl_path,config_path=config_path)
    
    woe_train_path = new_path +'/dataset_train_woed_' + target + hospital_name + '.csv'
    fp.process_woe_trans(data_path,rst_pkl_path,woe_train_path,config_path)
    
    info_value_dic = get_info_value_dic(feature_detail_path)
    
    print ('col length: ', len(col_name))
    # return 一个dic {feature name， infromation value}
    return info_value_dic

''' 如果 iv 值大于threshold， 取这个iv 值'''
def featue_selection_iv(info_value_dic, threshold, threshold_aa):
    # iv_list_touple (feature, iv value)
    iv_list_touple = [[i,info_value_dic[i]] for i in info_value_dic if (info_value_dic[i] > threshold)]
    # 删除 aminoAcid iv value that is > 1
    iv_list_touple = [i for i in iv_list_touple if ('aminoAcid' not in i[0] )or ('aminoAcid' in i[0] and i[1] < threshold_aa)]
    
    
    return iv_list_touple


''' 删除target 列 空值的行. Delete row based on nulls in certain columns '''
def target_missing(df, target):
    df_data1 = df.copy()
    if df[target].isnull().sum() != 0:
        df[target] = df[target].fillna('9999')
        # 找到 null 所在的行，并删除一整行
        null_list = df[(df[target] == '9999')].index.tolist()
   
        df_data1 = df.drop(null_list)
    return df_data1

''' 检查 hla_genes 异常'''
def check_hla_genes(df_data0):

    # 将 除了 hla_genes 的列相加，存成 df_sum column
    df_data_fill = df_data0.fillna(0)
    df_list = df_data_fill.columns.tolist()
    df_hla_col_list = df_list[df_list.index('hla_a'):df_list.index('hla_genes')]
    
    # 方便未来 debug: 比较原始数据和 df_col_sum, compare two columns， 有'DIFF!' 的就是计算异常
    # df_test = df_col_sum.where((df_col_sum.values == df_data0['hla_genes'].values), other='DIFF!')
    df_data0['hla_genes'] = df_data0[df_hla_col_list].sum(axis=1)
    return df_data0

'''删除无用列'''
def drop_unuseful_cols(using_ht_data, ht_data, df_data0, target):
    # 'kir_R_c1'-'D_R_A3A11'
    if using_ht_data:
        df_data_list = df_data0.columns.tolist()
        drop_cols = ['identity', 'ID', 'CG_GRD','DAGEGP', 'DISEASE', 'KIR_R_2DS4', 'KIR_D_AB', 'D_motif_B', 'R_motif_B', 'D_motif_C', 
                     'R_motif_C','D_ser_A', 'R_ser_A'] +  df_data_list[df_data_list.index('kir_R_c1'):df_data_list.index('RISK_level')]
        for i in drop_cols:
            if i in df_data0.columns:
                df_data0 = df_data0.drop(i, axis=1) 
    # is multi center data
    else:
        save_cols = [col for col in df_data0.columns if (col in ht_data.columns) or (col == target) or (col == 'D_KIR_Haplotype_C') or (col == 'R_KIR_Haplotype_C')]
        df_data0 = df_data0[save_cols]
       
    return df_data0

def change_content_to_underscore(df_data3, col_1, col_2):
    #'D_KIR_Haplotype_C' 'D_KIR_Haplotype_T'
    df_data3[col_1] = df_data3[col_1].fillna('none')
    df_data3[col_2] = df_data3[col_1].fillna('none')
    new_col_D = [i.replace("/", "") for i in df_data3[col_1] if i != 'none']
    new_col_R = [i.replace("/", "") for i in df_data3[col_2] if i != 'none']
    
    #df_data3 = pf.merge(df_data3, new_col_D,how='outer', on='ID')
    df_data = {col_1+'_new':new_col_D, col_2+'_new':new_col_R}
    df = pd.DataFrame.from_dict(df_data)
    df_data3 = df_data3.join(df)
    df_data3 = df_data3.drop(col_1, axis=1)
    df_data3 = df_data3.drop(col_2, axis=1)
    
    return df_data3

''' 不用存 null 的 one hot encoding'''    
def delete_one_hot_null(df_data3):
    # for debug only
    col_null_list = [col for col in df_data3.columns if ('*' in col) or ('?' in col)]
    col_not_null_list = [col for col in df_data3.columns if '*' not in col]
    df_data3 = df_data3[col_not_null_list]
    return df_data3, col_null_list

''' eg. 将非常相关的两个feature 删掉一个留一个'''
def feature_correlation(df_data5, threshold=0.95):
    
    # Create correlation matrix
    corr_matrix = df_data5.corr().abs()
    # Select upper triangle of correlation matrix
    upper = corr_matrix.where(np.triu(np.ones(corr_matrix.shape), k=1).astype(np.bool))
    # Find index of feature columns with correlation greater than 0.98
    to_drop = [column for column in upper.columns if any(upper[column] > threshold)]
    list_not_drop = [i for i in df_data5.columns if i not in to_drop] # list(set(df_data5.columns) - set(to_drop))
    
    # update df_data5 & iv_list_touple
    df_data5 =  df_data5[list_not_drop]
    #iv_list_touple = [i for i in iv_list_touple if i[0] in list_not_drop]
    return df_data5

def feature_combination(df_data, col1, col2, operation):
    #col1 = 'AGE'
    #col2 = 'R_DPB1_aminoAcid_p178_L'
    # add ==== AGE_add_R_DPB1_aminoAcid_p178_L
    
    if operation == 'add':
        df_data[col1 +'_'+ operation +'_'+ col2] = df_data[col1] + df_data[col2]
    elif operation == 'subtract':
        df_data[col1 +'_'+ operation +'_'+ col2] = df_data[col1] - df_data[col2]
    elif operation == 'multiply':
        df_data[col1 +'_'+ operation +'_'+ col2] = df_data[col1] * df_data[col2]
    else:
        df_data[col1 +'_'+ operation +'_'+ col2] = df_data[col1] / df_data[col2]
    
    return df_data

''' Step 0 - Step 5'''
def data_processing(hospital_name, using_ht_data, ht_data, df_data0, new_path, col_start, col_end, need_to_recalculate_iv, need_normalize,target_variables_col,target):
 
    print('program start now: ',datetime.now())
    
    ''' STEP 0: 检查数据 hla_genes 是否为 hla_a +...+ hla_dp 的和 && 删除无用列'''
    ''' Drop: Disease / KIR_R_2DS4 '''
    df_data0 = df_data0[df_data0.columns.tolist()[col_start:col_end]]
    df_data0 = check_hla_genes(df_data0)
    df_data0 = seperate_patient_donor_blood_type(df_data0)

    df_data0 = drop_unuseful_cols(using_ht_data, ht_data, df_data0, target)
    
    ''' STEP 1: 对于每个target 将target missing 的 row 给删除掉 && normalization ============= DATA 1 '''
    
    col_selected = [i for i in df_data0.columns if (i == target) or (i not in target_variables_col)]
    df_data0 = df_data0[col_selected]
    df_data1 = target_missing(df_data0, target)
    if need_normalize:
        df_data1= normalization(df_data1)
        
    print('STEP 1: 处理 missing target :', datetime.now())
    
    ''' STEP 2: 将missing rate 大于阈值的 feature column 删除 ============= DATA 2 '''
    # 筛除之后，检查是否有target 被 删除掉
    df_data2 = df_data1.copy()
    df_data2_list = missing_value_selection(df_data2, df_data2.columns, 0.6)
    
    # 生成 DATA 2 & 定义 Feature Variable
    df_data2 = df_data1[df_data2_list]
    
    # 删除非常有关联的特征
    df_data2 = feature_correlation(df_data2, threshold = 0.98)
    
    #features_variables_col = list(set(df_data2.columns) - set(target_variables_col))
    print('STEP 2: 处理 missing rate :', datetime.now())
    
    ''' OPTIONAL STEP 3: one hot encoding ============= DATA 3 '''
    # 先将 kir_Haplotype A/B -> A_B 方便 one hot encoding 命名
    df_data3 = df_data2.copy()
    df_data3 = change_content_to_underscore(df_data3, 'D_KIR_Haplotype_C', 'D_KIR_Haplotype_T')
    # 为了 one hot encoding, 先进行一个数据拆分 eg. 将 aminoacid 拆分成 aminoacid_L 
    cols_one_hot = [i for i in df_data3.columns if (('aminoAcid' in i) or (i == 'DIS') or (i == 'D_KIR_Haplotype_C_new') or (i == 'D_KIR_Haplotype_T_new') or (i == 'BLD_D_A_patient') or (i == 'BLD_D_A_donor' ))]
    df_data3.to_csv(path + 'data_before_one_hot_encoding'+ target + hospital_name +'.csv', )
    df_data3 = one_hot_encoding(df_data3, cols_one_hot)
    # delete one hot encoding for null 
    df_data3,col_null_list = delete_one_hot_null(df_data3)
    
    
    # 强强feature 联合
#    df_data3 = feature_combination(df_data3, 'D_A_aminoAcid_n20_M', 'D_A_aminoAcid_n19_A', 'add')
#    df_data3 = feature_combination(df_data3, 'D_A_aminoAcid_n20_M', 'D_A_aminoAcid_n19_A', 'subtract')
#    df_data3 = feature_combination(df_data3, 'D_A_aminoAcid_n20_M', 'D_A_aminoAcid_n19_A', 'multiply')
#    df_data3 = feature_combination(df_data3, 'D_A_aminoAcid_n20_M', 'D_A_aminoAcid_n19_A', 'devide')
#    
    # update feature
    features_variables_col = [i for i in df_data3.columns if i not in target_variables_col] #list(set(df_data5.columns) - set(target_variables_col))
    print('OPTIONAL STEP: one hot encoding: ', datetime.now())
    
    ''' STEP 4: 填充 缺失值 (可以用 data 2 / data 3) ============= DATA 4 '''
    # 分成 col_cat, col_num_cot, 和 col_num_dis 来进行填充
    # col_cat: 填充一个没有出现过的数字/ 字符串
    # col_num_cot: 填充均值
    # col_num_dis (HLA):填充均值
    df_data4 = df_data3.copy()
    col_cat, col_num_cot, col_num_dis = cat_cot_dis_data(df_data4, threshold =10)
    fill_num = 27
    df_data4 = fill_missing_value(df_data4,col_cat, col_num_cot, col_num_dis, fill_num)
    
    print('STEP 4: 填充 缺失值: ', datetime.now())
    
    ''' STEP 5: IV ============= DATA 5 '''
    df_data5 = df_data4.copy()
    info_value_dic = information_value(hospital_name,df_data5, new_path, col_num_cot, features_variables_col, target)
    csv_data5 = new_path + 'df_data5'+'_' + target + hospital_name +'.csv'
    df_data5.to_csv(csv_data5, index=False)
    return df_data1, df_data2, df_data3, df_data4, df_data5, info_value_dic



def data_processing_and_iv_value(hospital_name, using_ht_data, ht_data, df_data0, new_path, col_start, col_end, need_to_recalculate_iv, need_normalize,target_variables_col,target, is_woe_data):

    if (need_to_recalculate_iv):
        # df_data 0 - 5 is for debug only && 重新算 iv
        df_data1, df_data2, df_data3, df_data4, df_data5, info_value_dict = data_processing(hospital_name, using_ht_data, ht_data, df_data0, new_path, col_start, col_end, 
                                                                                            need_to_recalculate_iv, need_normalize,
                                                                                            target_variables_col,target)
    else:
        # 因为 df iv 已经算过一遍了，直接从这里开始读即可
        # 读取 woe 值
        if is_woe_data:
            df_data5 = pd.read_csv(new_path + 'dataset_train_woed_' + target + hospital_name +'.csv')
            df_data5 = df_data5[df_data5.columns.tolist()[1:]]
            info_value_dict = get_info_value_dic(new_path + 'features_detail_' + target + hospital_name + '.csv')
        # 读取经过数据清洗过的 data
        else:
            info_value_dict = get_info_value_dic(new_path + 'features_detail_' + target + hospital_name + '.csv')
        
    df_data6 = df_data5.copy()
    return df_data6, info_value_dict

''' 归一化 '''
def normalization(df):
    # for debug only
    #df=df[df.columns.tolist()[0:3]]
    df=[(df[i]-df[i].min())/(df[i].max()-df[i].min()) for i in df.columns]
    return df 


''' target could be "target" for woe data, or "eg.CGVHD" for original data '''
def feature_selection(df_data6, info_value_dict,is_woe_data, new_path, target):
    # 进行第一步筛选，只保留特征值(iv) 大于 threshod 的feature, threshod 初步定在0.03
    # 删除 aminoAcid iv value that is > 1 [0.3 1.5]
    threshold = 0.01 
    threshold_aa_max = 1.0
    version = '3'
    target_name = 'target' if is_woe_data else target

    
    # two csv ======  feature count & analysis
    # get the col for all the data 
    csv_feature_names = df_data6.columns.values.tolist()
    
    df_data = {'Features':csv_feature_names}
    df_feature_summary = pd.DataFrame.from_dict(df_data)

    #first_col = df_data6.loc[0](axis=1)
    
    #for threshold in np.linspace(0.15, 0.19, num=2):
        #for threshold_aa_max in np.linspace(1.5, 4, num=2):
    for threshold in np.linspace(0, 0.3, num=15):
    #if True:
        #if True:
        iv_list_touple = featue_selection_iv(info_value_dict, threshold, threshold_aa_max)
        print('STEP 5: IV : ', datetime.now())

        # 看看选择了哪些重要特征
        iv_list_touple = sorted(iv_list_touple, key = itemgetter(1),reverse=True)
        print('\t','threshold: ', threshold,'\t','threshold_aaMAX: ', threshold_aa_max,'\n\n')
        #print(iv_list_touple[0:11])
        
        ''' STEP 6: Scikit Learn feature selection ============= DATA 6 '''
        
        ## test
        #cols_selected = [col for col in df_data6.columns.tolist()]
        cols_selected = [col for col in df_data6.columns.tolist() if ('amino' not in col) or (col == 'target')]
        df_data6 = df_data6[cols_selected]
        #iv_list_touple = [l for l in iv_list_touple]
        iv_list_touple = [l for l in iv_list_touple if 'amino' not in l[0]]
        iv_list_touple = [l for l in iv_list_touple if l[1] < 1.0]
        print(len(iv_list_touple))
        
        df_train, df_test = train_test_split(df_data6, test_size=0.3, random_state=7)
        
#            if (woe_data):
        array_feature_train = df_train[[x[0] for x in iv_list_touple if x[0] != target_name]].values
        array_target_train = df_train[target_name].values
        
        array_feature_test = df_test[[x[0] for x in iv_list_touple if x[0] != target_name]].values
        array_target_test = df_test[target_name].values

#            else:
#                array_feature_train = df_data6[[x[0] for x in iv_list_touple if x[0] != 'CG_GRD']].values
#                array_target_train = df_data6['CGVHD'].values
#                
#                array_feature_test = df_test[[x[0] for x in iv_list_touple if x[0] != 'CG_GRD']].values
#                array_target_test = df_test['CGVHD'].values
#        
        
        # random forest
        estimators_dict = {'logreg': lm.LogisticRegression(),
                           'svc_linear': SVC(kernel="linear"),
                           'rf': RandomForestClassifier(n_estimators=500, min_samples_leaf=5, n_jobs=-1),
                           'xgb': xgb.XGBClassifier(max_depth = 8,n_estimators=1000
                                                    ,learning_rate=0.05,nthread=4
                                                    ,subsample=0.8,colsample_bytree=0.5
                                                    ,min_child_weight = 7)
                           }
        est = 'logreg'
        rfecv = RFECV(estimator=estimators_dict[est], step=1, cv=StratifiedKFold(10),
                  scoring='roc_auc')
        # 之前遇到的 bug： 不能将 str 变成 float。解决办法：先都 在step 3 one hot encoding，再来跑这里
        rfecv.fit(array_feature_train, array_target_train)
        print("Optimal number of features : %d" % rfecv.n_features_)
        #print( rfecv.ranking_)
      
        array_predict = rfecv.predict(array_feature_test)
        auc = roc_auc_score(array_target_test, array_predict)
        print('auc_TEST: ', auc)
    
        # Plot number of features VS. cross-validation scores
        fig = plt.figure(figsize=(18, 12))
        plt.title(target+' (Train Dadaset)')
        plt.xlabel("Number of features selected")
        plt.ylabel("Cross validation score (nb of correct classifications)")
        plt.plot(range(1, len(rfecv.grid_scores_) + 1), rfecv.grid_scores_)
        fig_path_part0 = new_path + target +'/feature_selection_pics/' +  target + '_v_' + version + '_' + est + '_maxAUC_' + str("{0:.4f}".format(max(rfecv.grid_scores_))) 
        fig_path_part1 =  '_finalAUC_' + str("{0:.4f}".format(rfecv.grid_scores_[-1])) + '_auc_TEST_' + str("{0:.4f}".format(auc))
        fig_path_part2 =  '_threshold_' +  str("{0:.4f}".format(threshold)) + '_threshold_aaMAX_' + str("{0:.4f}".format(threshold_aa_max))  + '_feature_' + str(rfecv.n_features_) +'.pdf'
        fig.savefig(  fig_path_part0 + fig_path_part1 + fig_path_part2 , dpi=fig.dpi)
    
        #plt.show()
        print('\n',fig_path_part1 + fig_path_part2)
    
        csv_thre_and_thre_aa_part1 = str("{0:.4f}".format(threshold)) + '_' + str("{0:.4f}".format(threshold_aa_max)) + '_' + "features_%d" % rfecv.n_features_
        csv_thre_and_thre_aa_part2 = '_maxAUC_' + str("{0:.4f}".format(max(rfecv.grid_scores_)))+ str("_{0:.4f}".format(auc))
        df_feature_summary[csv_thre_and_thre_aa_part1 + csv_thre_and_thre_aa_part2] = 0
        for i in range(0, len(iv_list_touple)):
            if rfecv.ranking_[i] == 1:
                df_feature_summary.at[csv_feature_names.index(iv_list_touple[i][0]), csv_thre_and_thre_aa_part1 + csv_thre_and_thre_aa_part2] = 1
    
    
    # store all feature to a csv file
    df_feature_summary.to_csv(new_path+'feature_selection_summary.csv', index=False, encoding= 'utf-8')
    return df_data6, df_feature_summary

''' 建造一个best model '''
def grid_search(original_model):
    
    # Create regularization penalty space
    penalty = ['l1', 'l2'] 
    # Create regularization hyperparameter space
    C = np.logspace(0, 4, 6, 10, 30)
    class_weight = [{0:0.8,1:0.2},{0:0.5,1:0.5},{0:0.4,1:0.6},{0:0.3,1:0.7},{0:0.2,1:0.8},{0:0.1,1:0.9}, 'balanced']
    # Create hyperparameter options
    hyperparameters = dict(C=C, penalty=penalty, class_weight =class_weight)
    # Create grid search using 5-fold cross validation
    clf = GridSearchCV(original_model, hyperparameters, cv=5, verbose=0)
    
    return clf






#
#    param_test1 = {
#     'max_depth':range(3,10,2),
#     'min_child_weight':range(1,6,2)
#    }
#    
#    gsearch1 = GridSearchCV(estimator = xgb.XGBClassifier( learning_rate =0.1, n_estimators=140, max_depth=5,
#     min_child_weight=1, gamma=0, subsample=0.8, colsample_bytree=0.8,
#     objective= 'binary:logistic', nthread=4, scale_pos_weight=1, seed=27), 
#     param_grid = param_test1, scoring='roc_auc',n_jobs=4,iid=False, cv=5)
#                
#    gsearch1.fit(train[predictors],train[target])
#    gsearch1.grid_scores_, gsearch1.best_params_, gsearch1.best_score_
#    
    
    
#    if str_model == 'best_model':
#        print("Best: %f using %s" % (model.best_score_, model.best_params_))
#    
#        print('************', model.best_estimator_)
#        model = model.best_estimator_
#        print(model)

'''用state 变换来评价模型，对不同的state的模型效果取平均，方差可以得到模型的靠谱效果'''
def check_model_avg_performance_using_state(hospital_name, multi_center_path, model, str_model, X, y):
#for model in [original_model, best_model]:
    #document.add_paragraph ('****************************')
#    if model == original_model:
#        str_model = 'model_' 
#        count = 1
#    else: 
#        str_model = 'best_model_'
#        count = 2
    document = Document()
    document.add_heading('Model Coefficient Summary', 0)    
    
    variance = []
    mean_value = []
    for state in range(1,10):
        x_train, x_test, y_train, y_test  = train_test_split(X, y, test_size=0.3, random_state=state)#default state 6                       
        model = model.fit(x_train, y_train)    
        
        auc_train = cross_val_score(model, x_train, y_train, cv=10, scoring='roc_auc')
        str_auc_train = str_model + "auc_train_%0.4f_%0.2f" % (auc_train.mean(), auc_train.std() * 2)
    #            str_auc_train = "auc_train: %0.2f (+/- %0.2f)" % (auc_train.mean(), auc_train.std() * 2)
        #document.add_paragraph(str_auc_train)    
    
        # auc train
        probs_test = model.predict_proba(x_test)
        preds_test = probs_test[:,1]
        fpr_test, tpr_test, threshold = metrics.roc_curve(y_test, preds_test)
        auc_test = metrics.auc(fpr_test, tpr_test)
        str_auc_test = "auc_test_%0.4f" % auc_test
        #document.add_paragraph(str_auc_test)                        
    
        # make class predictions for the testing set
        y_pred = model.predict(x_train)
        acc_train = metrics.accuracy_score(y_train, y_pred)  
        str_acc_train = "acc_train_%0.4f_%0.2f" % (acc_train.mean(), acc_train.std() * 2)
    #            str_acc_test = "acc_test: %0.2f (+/- %0.2f)" % (acc_test.mean(), acc_test.std() * 2)
        #document.add_paragraph(str_acc_train)    
    
        # calculate the fpr and tpr for all thresholds of the classification
        probs = model.predict(x_test)
        acc_test = metrics.accuracy_score(y_test, probs)
        str_acc_test = "acc_test_%0.4f" % acc_test
        
        variance.extend((auc_train.mean(), auc_test, acc_train.mean(), acc_test))
        mean_value.extend((auc_train.mean(), auc_test, acc_train.mean(), acc_test))
        
        if str_model == 'best_model':
            document.add_paragraph('************')
            document.add_paragraph("Best: %f using %s" % (model.best_score_, model.best_params_))
            document.add_paragraph(str(model.best_estimator_))
            
            #print("Best: %f using %s" % (model.best_score_, model.best_params_))
        
            #print('************', model.best_estimator_)
            #model = model.best_estimator_
            #print(model)
        
            
        
        ''' draw roc curve''' 
        fig = plt.figure(figsize=(18, 12))
        plt.title('Receiver Operating Characteristic')
        plt.plot(fpr_test, tpr_test, 'b', label = 'AUC = %0.2f' % auc_test)
        plt.legend(loc = 'lower right')
        plt.plot([0, 1], [0, 1],'r--')
        plt.xlim([0, 1])
        plt.ylim([0, 1])
        plt.ylabel('True Positive Rate')
        plt.xlabel('False Positive Rate')
        #plt.show()
    
        # + '_'.join(cols)
        fig_path_part1 = new_path + target + '/pics/' + model_type + '_' + target + '_roc_curve_' + 'state_' + str(state) + '_'
        fig_path_part2 = '_'.join([str_auc_train, str_auc_test, str_acc_train, str_acc_test]) +'.pdf'
        fig.savefig(fig_path_part1 + fig_path_part2 , dpi=fig.dpi)
        print(fig_path_part1 + fig_path_part2)
        
        #draw_roc_curve(new_path, fpr_test, tpr_test, auc_test, model_type, target, state, str_auc_train, str_auc_test, str_acc_train, str_acc_test)
    
    #document.add_paragraph(str_acc_test) 
    print(str_model, 'variance is: ', np.var(variance))  
    print(str_model, 'mean is: ',mean(mean_value))
    print(str_model, 'auc test: ',mean(mean_value[1::4]))
    print(str_model, 'auc train: ',mean(mean_value[0::4]))
    print(str_model, 'acc test: ',mean(mean_value[3::4]))
    print(str_model, 'acc train: ',mean(mean_value[2::4]))
    #document.save(multi_center_path +'model_estimator_from_grid_search' + '_' + target + hospital_name + '.docx')           
    return model
#def draw_roc_curve(new_path, fpr_test, tpr_test, auc_test, model_type, target, state, str_auc_train, str_auc_test, str_acc_train, str_acc_test):
        

def model_and_roc_curve(hospital_name, multi_center_path, df_data6, cols, model, model_type, X, y):
#    result = []
#    if model_type == 'logistic':
#        logit_model=sm.Logit(y,X)
#        result=logit_model.fit()
#        print(result.summary())
#    elif model_type == 'xgb':
#        pass
    
    print ('STEP 6: start model selection')
    #list_models = []
    #important_states = [3, 7, 10, 12, 14]
    #document.add_heading('**************', level=1)
#    for state in important_states: # range(1,20):
#    if True:
#        state = 7
        #document.add_paragraph('state: '+ str(state))
        # making an instance of model
    original_model = model #C=2.7825594022071245
    original_model = check_model_avg_performance_using_state(hospital_name, multi_center_path, original_model,'model', X, y)  
    best_model = original_model
    
    #clf = grid_search(original_model)
    #best_model = check_model_avg_performance_using_state(hospital_name, multi_center_path, clf,'best_model', X, y).get_params()['estimator']
     
    #print_coef(model, target, X, y, 'model')
    #print_coef(best_model, target, X, y, 'best_model')
    

    
    return original_model, best_model


def print_coef(model, target, X, y, str_model):
    
    document = Document()
    document.add_heading('Model Coefficient Summary', 0)    
    
    
    if str_model == 'model':
        logit_model = model
    else: 
        logit_model = model.get_params()['estimator']
    logit_model.fit(X, y)
    document.add_paragraph('---------------------------')
    document.add_paragraph('For target:' + target + '\n')
#    for j in range(len(logit_model.coef_[0])):
#        print('   ', cols[j], '=',logit_model.coef_[0][j])
        
        
        
        
        #document.add_paragraph('   '+ cols[j]+ '=' + str(logit_model.coef_[0][j]))
    #document.save(new_path+'model_coefficient_summary.docx')   

        
#        print('For target:', target)
#        for j in range(len(logit_model.coef_[0])):
#            print('   ', cols[j], '=',logit_model.coef_[0][j])
#     
#    logit_model = lm.LogisticRegression(C=2.7825594022071245, class_weight=None, dual=False,
#          fit_intercept=True, intercept_scaling=1, max_iter=100,
#          multi_class='ovr', n_jobs=1, penalty='l2', random_state=None,
#          solver='liblinear', tol=0.0001, verbose=0, warm_start=False)
#    
    
    
    return 

def HTYDSH_df_woe(df_data6_list, multicenter_path, col_num_cot, features_variables_col):
    df_total = pd.concat(df_data6_list).fillna(0)
    info_value_dic = information_value('_HTYDSH_',df_total, multicenter_path, col_num_cot, features_variables_col, 'target')
    csv_data5 = new_path + 'df_data5'+'_' + target + '_HTYDSH_' +'.csv'
    df_total.to_csv(csv_data5, index=False)
    return df_total

''' fit 一个模型，并用这个模型来预测剩下的其他医院的数据'''
def use_one_model_to_perdict_the_other():
    
     #建立逻辑回归模型
    cols=['AGE','DAGE','DSEX', 'D_A3_new', 'kir_D_2DL1', 'DSEX','D_C1C2', 'kir_D_2DS1'
          ,'D_2DS3_2DS5', 'R_Missing_A3', 'R_Missing_A3A11','R_Missing_Bw4_Licensed', 'kir_D_2DS5', 'PIRCHE_II_num', 'D_A11_new', 'D_A3A11_new'
          , 'PIRCHE_I_num', 'R_C1C2', 'R_C1C2_D_C1C1','Effect_Haplotype_D_BX_C1C2', 'hla_genes'] 
    
    
    for df in df_data6_list:
        for col in cols:
            if col not in df.columns:
                df[col] = [0] * df.shape[0] 
    
    # 
    #cols = [k for k,v in info_value_dict.items() if v > 0.02 and v < 2.2 and 'amino' not in k]
    
    model = lm.LogisticRegression(C=20.0, class_weight={0: 0.5, 1: 0.5},  dual=False,
          fit_intercept=True, intercept_scaling=1, max_iter=150,
          multi_class='ovr', n_jobs=1, penalty='l2', random_state=None,
          solver='liblinear', tol=0.0001, verbose=0, warm_start=False)
    
    
    #model = XGBClassifier()
    
    
    
    df_total = pd.concat(df_data6_list).fillna(0)
    X = df_total[cols]
    y = df_total['target']
    original_model, best_model = model_and_roc_curve('_HTYDSH_',path, df_total, cols, model, model_type, X, y)
    
    
    
    count = 0
    str_name = ''    
    for df_data6 in df_data6_list:
        
        print('\n',hospital_name_list[count])

        X = df_data6[cols]
        y = df_data6['target']
        probs = original_model.predict(X)
        acc_test = metrics.accuracy_score(y, probs)
        str_acc_test = "acc_test_%0.4f" % acc_test
        print(str_acc_test)
        
        probs_test = original_model.predict_proba(X)
        preds_test = probs_test[:,1]
        fpr_test, tpr_test, threshold = metrics.roc_curve(y, preds_test)
        auc_test = metrics.auc(fpr_test, tpr_test)
        str_auc_test = "auc_test_%0.4f" % auc_test
        print(str_auc_test)
        
        str_name += hospital_name_list[count] + str_acc_test + str_auc_test
#        original_model, best_model = model_and_roc_curve(hospital_name_list[count],path, df_data6, cols, model, model_type, X, y)
        count += 1
    
    
    
    # save the model to disk
    filename = '_' + str_name + '_' +'.pkl'
    pickle.dump(model, open(multicenter_path + filename, 'wb'))
    return
    


if __name__ == '__main__':
    data_file_list = ['data/HT_20171031_MJZ1_Extension(20180312).xlsx' # 
                     ,'data/HT_data_merge_aminoAcid_raw.xlsx' #
                     ,'data/HT_data_Pirche.xlsx' # 
                     ]
    
    old_path = '/Users/DihuaDuolan/Desktop/BFR/task1To4/'
    # 因为在step5 要保存到 config，woe 以及很多的新文件，要 pass new path 给function
    new_path = '/Users/DihuaDuolan/Desktop/BFR/task5To8/'
    multicenter_path = '/Users/DihuaDuolan/Desktop/BFR/multi_center/'
    data_excel = 'output_combined.xlsx'
    multi_center_data_excel = 'HT_YD_SH_data_20180822.xlsx'
    # if False: is using multi-center data
    using_ht_data = False
    df_data = []

    woe_config_PATH_dict = {'HT':'/Users/DihuaDuolan/Desktop/BFR/multi_center/config_csv_CGVHD_HT_.csv'
                            ,'YD':'/Users/DihuaDuolan/Desktop/BFR/multi_center/config_csv_CGVHD_YD_.csv'
                            ,'SH':'/Users/DihuaDuolan/Desktop/BFR/multi_center/config_csv_CGVHD_SH_.csv'
                            ,'HTYDSH':'/Users/DihuaDuolan/Desktop/BFR/multi_center/config_csv_CGVHD_HTYDSH_.csv'}
    center_train = 'HTYDSH'
    
    

    df_config = pd.read_csv(woe_config_PATH_dict[center_train])
    cfg_cols = df_config['var_name'].tolist()
    

    
    if using_ht_data:
        df_data_original = pd.read_excel(old_path + data_excel, 'Sheet1')
        ht_data = df_data_original
        path = new_path
        df_data.append(df_data_original)
    else:
        df_data_original = pd.read_excel(multicenter_path + multi_center_data_excel, 'Sheet1') # multi center data
        ht_data = pd.read_csv(multicenter_path + 'original_data_woe_CGVHD_HT_.csv')
        path = multicenter_path
        
        df_data_original = df_data_original.rename(columns={'CGVHD': 'target'
                                                            ,'D_KIR_Haplotype_C':'D_KIR_Haplotype_C_new'
                                                            ,'D_KIR_Haplotype_T':'D_KIR_Haplotype_T_new'})
        
        df_data_original['D_KIR_Haplotype_C_new'] = df_data_original['D_KIR_Haplotype_C_new'].astype(str).apply(lambda x: ''.join(x.split('/')))
        df_data_original['D_KIR_Haplotype_T_new'] = df_data_original['D_KIR_Haplotype_T_new'].astype(str).apply(lambda x: ''.join(x.split('/')))
        
        df_data_original = one_hot_encoding(df_data_original, ['D_KIR_Haplotype_C_new','D_KIR_Haplotype_T_new'])

        df_data_original_copy = df_data_original.copy()
        
        df_data_original = df_data_original[cfg_cols]
        
        df_data_original['Center'] = df_data_original_copy['Center']
        df_data_original = df_data_original[~df_data_original['target'].isnull()]
        
        data_HT = df_data_original.loc[df_data_original['Center'] == 'Hang Tian'].drop(['Center'],axis=1).reset_index(drop=True)
        data_YD = df_data_original.loc[df_data_original['Center'] == 'Yan Da'].drop(['Center'],axis=1).reset_index(drop=True)
        data_SH = df_data_original.loc[df_data_original['Center'] == 'Shang Hai'].drop(['Center'],axis=1).reset_index(drop=True)
        data_HTYDSH = df_data_original.drop(['Center'],axis=1).reset_index(drop=True)
        df_data.extend((data_HT, data_YD, data_SH))

        
    
    # 保存一份最原始的数据
    df_data0 = df_data_original.copy()
    
    '''特征选择之前的 数据处理 '''
    col_start = 0
    col_end = -1
    need_to_recalculate_iv = False
    need_normalize = False
    is_woe_data = True
   
    model_type = 'logistic'
    
    
    # 定义 Target Variable, 已经算过：CGVHD， AGVHD
    target_variables_col = ['AG24', 'AG34', 'AGVHD', 'CGVHD', 'DEAD', 'REL', 
                            'TRM', 'DFS', 'CMV_y', 'EBV_y', 'CYSTITIS']
    target = target_variables_col[3]
    
    # 计算每个医院单独的 woe 值
    need_to_recalculate_iv = True
    data_processing_and_iv_value('HT', using_ht_data, ht_data, df_data0, new_path, col_start, col_end, need_to_recalculate_iv, need_normalize,target_variables_col,target, is_woe_data)
#    
#   
##    
##    if woe_data:
##        target = 'target'
##        # 将新读取的数据 改成 ’target‘
##        data_HT = data_HT.rename(columns={target_variables_col[3]: 'target'})
##        data_YD = data_HT.rename(columns={target_variables_col[3]: 'target'})
##        data_SH = data_HT.rename(columns={target_variables_col[3]: 'target'})
##        df_data.extend((data_HT, data_YD, data_SH))
##
##
##    else:
##        target = target_variables_col[3]
##    
#    
#    # seperate three
#    df_per_model = []
#    df_data6_list = []
#    hospital_name_list = ['_HT_', '_YD_', '_SH_', '_HTYDSH_']
#    
#    
#    woe_rule_PATH_dict = {'HT':'/Users/DihuaDuolan/Desktop/BFR/multi_center/woe_rule_CGVHD_HT_.pkl'
#                            ,'YD':'/Users/DihuaDuolan/Desktop/BFR/multi_center/woe_rule_CGVHD_YD_.pkl'
#                            ,'SH':'/Users/DihuaDuolan/Desktop/BFR/multi_center/woe_rule_CGVHD_SH_.pkl'
#                            ,'HTYDSH':'/Users/DihuaDuolan/Desktop/BFR/multi_center/woe_rule_target_HTYDSH_.pkl'}
#
#
#
#    data_preWOE_path_dict = {'HT':'/Users/DihuaDuolan/Desktop/BFR/multi_center/data_preWOE_HT.csv'
#                        ,'YD':'/Users/DihuaDuolan/Desktop/BFR/multi_center/data_preWOE_YD.csv'
#                        ,'SH':'/Users/DihuaDuolan/Desktop/BFR/multi_center/data_preWOE_SH.csv'
#                        ,'HTYDSH':'/Users/DihuaDuolan/Desktop/BFR/multi_center/data_preWOE_HTYDSH.csv'
#                        }
#    data_afterWOE_path_dict = {'HT':'/Users/DihuaDuolan/Desktop/BFR/multi_center/data_afterWOE_HT.csv'
#                        ,'YD':'/Users/DihuaDuolan/Desktop/BFR/multi_center/data_afterWOE_YD.csv'
#                        ,'SH':'/Users/DihuaDuolan/Desktop/BFR/multi_center/data_afterWOE_SH.csv'
#                        ,'HTYDSH':'/Users/DihuaDuolan/Desktop/BFR/multi_center/data_afterWOE_HTYDSH.csv'
#                        }
#
#    data_HT.to_csv(data_preWOE_path_dict['HT'],index=False)
#    data_YD.to_csv(data_preWOE_path_dict['YD'],index=False)
#    data_SH.to_csv(data_preWOE_path_dict['SH'],index=False)
#    data_HTYDSH.to_csv(data_preWOE_path_dict['HTYDSH'],index=False)
#    
#    
#
#    
#
#
#### woe individualy
#    center_dict = {'HT':'Hang Tian'
#                   ,'YD':'Yan Da'
#                   ,'SH':'Shang Hai'
#                   }    
#    for center_predict in data_preWOE_path_dict.keys():#[center for center in data_preWOE_path_dict.keys() if center != center_train]:
#        center_train_ = center_predict
#        df_config = pd.read_csv(woe_config_PATH_dict[center_train_])
#        cfg_cols = df_config['var_name'].tolist()
#        df_data_original = df_data_original_copy[cfg_cols]
#        
#        df_data_original['Center'] = df_data_original_copy['Center']
#        df_data_original = df_data_original[~df_data_original['target'].isnull()]
#        
#        if center_train_ != 'HTYDSH':
#            data_center = df_data_original.loc[df_data_original['Center'] == center_dict[center_train_]].drop(['Center'],axis=1).reset_index(drop=True)
#            data_center.to_csv(data_preWOE_path_dict[center_train_],index=False)
#        else:
#            data_center = df_data_original.drop(['Center'],axis=1).reset_index(drop=True)
#            data_center.to_csv(data_preWOE_path_dict[center_train_],index=False)
#
#
#        fp.process_woe_trans(data_preWOE_path_dict[center_predict],woe_rule_PATH_dict[center_train_]
#                        ,data_afterWOE_path_dict[center_predict],woe_config_PATH_dict[center_train_])
#
###### woe tranformed from center_train
##    for center_predict in data_preWOE_path_dict.keys():#[center for center in data_preWOE_path_dict.keys() if center != center_train]:
##        fp.process_woe_trans(data_preWOE_path_dict[center_predict],woe_rule_PATH_dict[center_train]
##                        ,data_afterWOE_path_dict[center_predict],woe_config_PATH_dict[center_train])
##        
##        
#    
##    for data_idx in range(0, len(df_data)):
##        df_data0 = df_data[data_idx]
##        df_data6, info_value_dict = data_processing_and_iv_value(hospital_name_list[data_idx], using_ht_data, ht_data, df_data0, path, col_start, 
##                                                                                                     col_end, need_to_recalculate_iv,need_normalize, target_variables_col,target,is_woe_data)
##        df_data6_list.append(df_data6)
##     
#    # for debug    
#    data_HT_after_data_process = pd.read_csv(data_afterWOE_path_dict['HT'])
#    data_YD_after_data_process = pd.read_csv(data_afterWOE_path_dict['YD'])
#    data_SH_after_data_process = pd.read_csv(data_afterWOE_path_dict['SH'])
#    
#    df_data6_list = [data_HT_after_data_process,data_YD_after_data_process,data_SH_after_data_process]
#
#    # 把全部data 求一个woe值
#    #df_total = HTYDSH_df_woe(df_data6_list, multicenter_path, col_num_cot, features_variables_col)
#    
#    #df_HTYDSH_woe = pd.read_csv(multicenter_path + 'dataset_train_woed_target_HTYDSH_.csv')
#    
#    
#    # 对于center_train 做feature selection
#    df_center_train_woe = pd.read_csv(data_afterWOE_path_dict[center_train]).fillna(0).replace('missing',0)#replace({'missing': 0})
#    df_center_train_woe = df_center_train_woe[df_center_train_woe.columns.tolist()[1:]]
#    center_train_info_value_list_tuple = get_info_value_dic(multicenter_path + 'features_detail_' + target + '_' + center_train + '_.csv')
#    center_train_info_value_list_dict = dict(center_train_info_value_list_tuple)
#    
#    df_feature_selection, df_feature_summary = feature_selection(df_center_train_woe, center_train_info_value_list_dict,is_woe_data,multicenter_path, target)
#    # 计算最常被用的 feature
#    df_feature_summary['feature_usage_count'] = df_feature_summary[df_feature_summary.columns].sum(axis=1)
#    
#    feature_usage_count = df_feature_summary.nlargest(15, 'feature_usage_count')['Features'].tolist()
#
#    
#    print(feature_usage_count)
#    
#    #建立逻辑回归模型
#    
#    
#    cols=['PIRCHE_I_num', 'D_A3_new', 'D_2DS3_2DS5', 'R_Missing_Bw4_Licensed', 'R_Missing_A3'
#          ,'hla_genes', 'DAGE', 'Effect_Haplotype_D_BX_C1C2', 'R_Missing_A3A11', 'kir_D_2DL1',
#          'PIRCHE_II_num', 'D_A3A11_new']
#          
##    cols=['R_A11_new', 'R_A3A11_new', 'R_AML', 'D_KIR_Haplotype_C_new_AA'
##          , 'D_KIR_Haplotype_T_new_AA', 'hla_dp', 'PIRCHE_II_num', 'D_A3A11_new'
##          , 'PIRCHE_I_num', 'R_ALL', 'hla_dq', 'R_Missing_C2', 'kir_R_bw4_new', 'hla_a', 'hla_b']
#          
##          'kir_D_2DS1', 'kir_D_2DS5', 'kir_D_2DL4', 'kir_D_2DL1',
##          , 'R_Missing_A3', 'R_Missing_A3A11', 'PIRCHE_II_num','PIRCHE_I_num'
##          , 'Effect_Haplotype_D_BX_C1C2', 'hla_genes'] 
##    
#    
#    for df in df_data6_list:
#        for col in cols:
#            if col not in df.columns:
#                df[col] = [0] * df.shape[0] 
#
#    
#    
#    # 
#    #cols = [k for k,v in info_value_dict.items() if v > 0.02 and v < 2.2 and 'amino' not in k]
#    
#    model = lm.LogisticRegression(C=0.7, class_weight={0: 0.5, 1: 0.5},  dual=False,
#          fit_intercept=True, intercept_scaling=1, max_iter=150,
#          multi_class='ovr', n_jobs=1, penalty='l2', random_state=None,
#          solver='liblinear', tol=0.0005, verbose=0, warm_start=False)
#    
#    df_total = pd.concat(df_data6_list).fillna(0).replace('missing', 0)
#    
#    data_afterWOE_dict = {'HT':data_HT_after_data_process
#                          ,'YD':data_YD_after_data_process
#                          ,'SH':data_SH_after_data_process
#                          ,'HTYDSH':df_total}
#    
#    df_temp = data_afterWOE_dict[center_train]
#    cols_ = [col for col in cols if col in df_temp.columns]
#    cols = cols_
#    
#    
#    X = data_afterWOE_dict[center_train][cols].fillna(0)#df_total[cols]
#    y = data_afterWOE_dict[center_train]['target']#df_total['target']
#    original_model, best_model = model_and_roc_curve(center_train.join(['_','_']),path, data_afterWOE_dict[center_train], cols, model, model_type, X, y)
#    
#    count = 0
#    str_name = 'train_' + center_train + '_predict_'  
#    df_data6_list.append(df_total)
#    for df_data6 in df_data6_list:
#        
#        print('\n',hospital_name_list[count])
#
#        X = df_data6[cols].fillna(df_data6[cols].mean()).fillna(0)
#        y = df_data6['target']
#        probs = original_model.predict(X)
#        acc_test = metrics.accuracy_score(y, probs)
#        str_acc_test = "_acc_%0.4f" % acc_test
#        print(str_acc_test)
#        
#        probs_test = original_model.predict_proba(X)
#        preds_test = probs_test[:,1]
#        fpr_test, tpr_test, threshold = metrics.roc_curve(y, preds_test)
#        auc_test = metrics.auc(fpr_test, tpr_test)
#        str_auc_test = "_auc_%0.4f" % auc_test
#        print(str_auc_test)
#        
#        str_name += hospital_name_list[count] + str_acc_test + str_auc_test
##        original_model, best_model = model_and_roc_curve(hospital_name_list[count],path, df_data6, cols, model, model_type, X, y)
#        count += 1
#    
#    
#    
#    # save the model to disk
#    filename = '' + str_name + '_' +'.pkl'
#    pickle.dump(model, open(multicenter_path + filename, 'wb'))
# 
#    
#    
#    # first ten features
#    data_dia = y
#    data = X
#    data_n_2 = (data - data.mean()) / (data.std())              # standardization
#    data = pd.concat([y,data_n_2.iloc[:,0:10]],axis=1)
#    data = pd.melt(data,id_vars="target",
#                        var_name="features",
#                        value_name='value')
#    fig = plt.figure(figsize=(18, 12))
#    #plt.figure(figsize=(10,10))
#    
#    sns.violinplot(x="features", y="value", hue="target", data=data,split=True, inner="quart")
#    plt.xticks(rotation=90)
#    fig.savefig(multicenter_path + 'model_feature.pdf', dpi=fig.dpi)
#    
    
#    print('original model is: ')
#    print(original_model)
#    print('best model is: ')
#    print(best_model)
#    
    
    '''  '''
    #print_coef(model, target, X, y, list_models)
    
    
    #print(result.summary())
    #print(auc_score_str)
    #print(result.summary())
    #print(model)
    
    # try xgb model, first: with one hot encoding, with fillNAN
  
    
    # try xgb model, second: without one hot encoding, without fillNAN

    
    
    
#    
    
#    # Feature Combination
#    df_data6 = feature_combination(df_data2, 'AGE', 'DAGE', 'add')
#    df_data6 = feature_combination(df_data2, 'AGE', 'DAGE', 'subtract')
#    df_data6 = feature_combination(df_data2, 'PIRCHE_I_num', 'PIRCHE_II_num', 'add')
#    df_data6 = feature_combination(df_data2, 'PIRCHE_I_num', 'PIRCHE_II_num', 'subtract')
#    
    
    
   
